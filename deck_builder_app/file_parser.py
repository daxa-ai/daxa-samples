"""Reusable file -> plain-text parsing helpers (no Streamlit dependency).

Framework-agnostic on purpose: any caller (Streamlit app, CLI script, API
endpoint) can pass raw bytes in and get an LLM-friendly text blob out.

Supports .xlsx/.xlsm (spreadsheets) and .docx (Word documents) today; use
parse_file_to_text() to dispatch on filename extension automatically.
"""
import io
import zipfile
from typing import List

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from openpyxl import load_workbook
from openpyxl.utils.exceptions import InvalidFileException

MAX_CHARS = 8000
MAX_ROWS_PER_SHEET = 500

SUPPORTED_EXTENSIONS = (".xlsx", ".xlsm", ".docx")


class FileParsingError(Exception):
    """Raised when input bytes cannot be parsed as a supported file type."""


def is_supported_file(filename: str) -> bool:
    """Return True if filename's extension has a parser (.xlsx, .xlsm, .docx)."""
    return filename.lower().endswith(SUPPORTED_EXTENSIONS)


def parse_xlsx_to_text(file_bytes: bytes, filename: str = "", max_chars: int = MAX_CHARS) -> str:
    """Parse .xlsx bytes into a plain-text, LLM-friendly representation.

    One "Sheet: <name>" section per worksheet, followed by its rows (cell
    values joined with " | "), one row per line. Result is truncated to
    `max_chars`.

    Args:
        file_bytes: Raw bytes of an .xlsx/.xlsm file.
        filename: Original filename, used only to make error messages clearer.
        max_chars: Hard cap on the returned string's length.

    Returns:
        Plain-text description of every sheet's contents (never None; an
        empty or all-blank workbook returns descriptive placeholder text
        rather than raising).

    Raises:
        FileParsingError: if file_bytes is not a readable Excel workbook
            (wrong format, corrupt file, etc.).
    """
    try:
        workbook = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    except (InvalidFileException, KeyError, OSError, zipfile.BadZipFile) as exc:
        raise FileParsingError(
            f"Could not read '{filename or 'uploaded file'}' as an Excel (.xlsx) file: {exc}"
        ) from exc

    sections: List[str] = []
    for sheet_name in workbook.sheetnames:
        ws = workbook[sheet_name]
        lines = [f"Sheet: {sheet_name}"]
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            if row_count >= MAX_ROWS_PER_SHEET:
                lines.append("...[remaining rows truncated]")
                break
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):  # skip fully-empty rows
                lines.append(" | ".join(cells))
            row_count += 1
        if row_count == 0:
            lines.append("(empty sheet)")
        sections.append("\n".join(lines))

    text = "\n\n".join(sections) if sections else "(workbook has no sheets)"
    if len(text) > max_chars:
        text = text[:max_chars] + "\n...[truncated]"
    return text


def parse_docx_to_text(file_bytes: bytes, filename: str = "", max_chars: int = MAX_CHARS) -> str:
    """Parse .docx bytes into a plain-text, LLM-friendly representation.

    Paragraph text is emitted in order; table rows are joined with " | ",
    one row per line, interleaved where they appear in the document. Result
    is truncated to `max_chars`.

    Args:
        file_bytes: Raw bytes of a .docx file.
        filename: Original filename, used only to make error messages clearer.
        max_chars: Hard cap on the returned string's length.

    Returns:
        Plain-text description of the document's contents (never None; an
        empty document returns descriptive placeholder text rather than
        raising).

    Raises:
        FileParsingError: if file_bytes is not a readable Word document
            (wrong format, corrupt file, etc.).
    """
    try:
        document = Document(io.BytesIO(file_bytes))
    except (PackageNotFoundError, KeyError, OSError, zipfile.BadZipFile) as exc:
        raise FileParsingError(
            f"Could not read '{filename or 'uploaded file'}' as a Word (.docx) file: {exc}"
        ) from exc

    lines: List[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    text = "\n".join(lines) if lines else "(empty document)"
    if len(text) > max_chars:
        text = text[:max_chars] + "\n...[truncated]"
    return text


def parse_file_to_text(file_bytes: bytes, filename: str, max_chars: int = MAX_CHARS) -> str:
    """Dispatch to the right parser based on filename's extension.

    Raises:
        FileParsingError: unsupported extension, or the file fails to parse.
    """
    lower = filename.lower()
    if lower.endswith((".xlsx", ".xlsm")):
        return parse_xlsx_to_text(file_bytes, filename=filename, max_chars=max_chars)
    if lower.endswith(".docx"):
        return parse_docx_to_text(file_bytes, filename=filename, max_chars=max_chars)
    raise FileParsingError(
        f"Unsupported file type for '{filename or 'uploaded file'}'. "
        f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}."
    )
